"""Text extraction from PDF, DOCX, and TXT resume files."""

import logging
import os
import docx
import pdfplumber

logger = logging.getLogger(__name__)


class UnsupportedFileFormatError(Exception):
    """Exception for unsupported file formats."""

    def __init__(self, file_path: str, extension: str):
        self.file_path = file_path
        self.extension = extension
        super().__init__(f"Unsupported file format '{extension}' for file: {file_path}")


class TextExtractionError(Exception):
    """Exception for text extraction failures."""

    def __init__(self, file_path: str, original_error: Exception):
        self.file_path = file_path
        self.original_error = original_error
        super().__init__(
            f"Failed to extract text from {file_path}: {str(original_error)}"
        )


class TextExtractor:
    """Extract text from PDF, DOCX, or TXT files."""

    SUPPORTED_FORMATS = {".pdf", ".docx", ".txt"}

    def __init__(self):
        """Initialize the TextExtractor."""
        logger.info("TextExtractor initialized")

    def is_supported(self, file_path: str) -> bool:
        """Check if file format is supported."""
        extension = os.path.splitext(file_path)[1].lower()
        return extension in self.SUPPORTED_FORMATS

    def extract(self, file_path: str) -> str:
        """Extract text from PDF, DOCX, or TXT file."""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = os.path.splitext(file_path)[1].lower()
        logger.info(f"Extracting text from {extension} file: {file_path}")

        try:
            if extension == ".pdf":
                text = self._extract_pdf(file_path)
            elif extension == ".docx":
                text = self._extract_docx(file_path)
            elif extension == ".txt":
                text = self._extract_txt(file_path)
            else:
                raise UnsupportedFileFormatError(file_path, extension)

            logger.info(
                f"Successfully extracted {len(text)} characters from {file_path}"
            )
            return text

        except (UnsupportedFileFormatError, FileNotFoundError):
            raise
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            raise TextExtractionError(file_path, e)

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pdfplumber."""
        text = ""
        with pdfplumber.open(file_path) as pdf:
            logger.debug(f"PDF has {len(pdf.pages)} pages")

            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                    logger.debug(
                        f"Extracted {len(page_text)} chars from page {page_num}"
                    )
                else:
                    logger.warning(
                        f"Page {page_num} is empty or could not be extracted"
                    )

        return text.strip()

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx."""
        document = docx.Document(file_path)
        logger.debug(f"DOCX has {len(document.paragraphs)} paragraphs")

        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        logger.debug(f"Extracted {len(text)} characters from DOCX")
        return text.strip()

    def _extract_txt(self, file_path: str) -> str:
        """Extract text from plain text file with encoding detection."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            logger.debug(f"Read {len(text)} characters from TXT (UTF-8)")
        except UnicodeDecodeError:
            logger.warning(f"UTF-8 decoding failed, trying latin-1 for {file_path}")
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()
            logger.debug(f"Read {len(text)} characters from TXT (latin-1)")

        return text.strip()
